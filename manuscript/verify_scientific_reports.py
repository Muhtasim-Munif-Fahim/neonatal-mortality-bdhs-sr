"""Structural submission gate for the Scientific Reports manuscript package."""
from __future__ import annotations

import json
import re
import subprocess
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "manuscript" / "scientific_reports"
MAIN = OUT / "manuscript_scientific_reports.docx"
SUPP = OUT / "supplementary_information.pdf"
TITLE = ("Temporal evaluation of neonatal mortality prediction and decomposition "
         "of mortality decline in Bangladesh, 2011–2022")


def words(text: str) -> int:
    return len(re.findall(r"\b[\w'-]+\b", text))


def xml_names(path: Path) -> set[str]:
    with zipfile.ZipFile(path) as zf:
        return set(zf.namelist())


def document_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


def cited_reference_numbers(text: str) -> set[int]:
    """Parse sequential Nature citations, including en-dash ranges."""
    cited: set[int] = set()
    for group in re.findall(r"\[([0-9, \-\u2013]+)\]", text):
        for item in group.split(","):
            item = item.strip()
            if not item:
                continue
            parts = re.split(r"[-\u2013]", item)
            if len(parts) == 2:
                start, end = map(int, parts)
                cited.update(range(start, end + 1))
            else:
                cited.add(int(item))
    return cited


def main() -> int:
    failures: list[str] = []
    blockers: list[str] = []
    if not MAIN.exists():
        print(f"FAIL: missing {MAIN}")
        return 1

    doc = Document(MAIN)
    text = document_text(doc)
    lines = [x.strip() for x in text.splitlines() if x.strip()]
    if not lines or lines[0] != TITLE:
        failures.append("exact title")

    match = re.search(r"Abstract\s+(.*?)\s+Keywords:", text, re.S | re.I)
    if not match or words(match.group(1)) > 195:
        failures.append("unstructured abstract <=195 words")
    kw = re.search(r"Keywords:\s*(.+)", text, re.I)
    if not kw or len([x for x in kw.group(1).split(";") if x.strip()]) > 6:
        failures.append("at most six keywords")

    required_order = ["Introduction", "Results", "Discussion", "Methods",
                      "Data Availability", "Code Availability", "References",
                      "Acknowledgements and Funding", "Author Contributions",
                      "Additional Information", "Figure Legends", "Tables"]
    heading_positions = {
        paragraph.text.strip(): index
        for index, paragraph in enumerate(doc.paragraphs)
        if paragraph.text.strip() in required_order
    }
    positions = [heading_positions.get(x, -1) for x in required_order]
    if any(x < 0 for x in positions) or positions != sorted(positions):
        failures.append("required declaration/section order")
    discussion = text[text.find("Discussion"):text.find("Methods")]
    if any(p.style.name.startswith("Heading 2") for p in doc.paragraphs
           if p.text and p.text in discussion):
        failures.append("Discussion must have no subheadings")

    legends = re.findall(r"^Figure\s+[1-5]\.\s+.+", text, re.M)
    table_titles = re.findall(r"^Table\s+[1-3]\.\s+.+", text, re.M)
    if len(legends) != 5 or len(table_titles) != 3:
        failures.append("exactly five figure legends and three main tables")
    if len(doc.tables) != 3:
        failures.append("exactly three editable main tables")

    names = xml_names(MAIN)
    embedded = [n for n in names if n.startswith("word/media/")]
    if embedded:
        failures.append("no embedded main figures")
    if "word/comments.xml" in names:
        failures.append("no comments")
    with zipfile.ZipFile(MAIN) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
        # Match revision elements, not table properties such as ``w:insideH``.
        if re.search(r"<w:(?:ins|del)(?:\s|>)", xml):
            failures.append("no tracked changes")
        if "w:lnNumType" not in xml:
            failures.append("line numbering")
        footer_xml = "".join(
            zf.read(n).decode("utf-8") for n in names if n.startswith("word/footer")
        )
        if "PAGE" not in footer_xml:
            failures.append("page numbering")

    if MAIN.stat().st_size >= 3_000_000:
        failures.append("main DOCX below 3 MB")
    if not SUPP.exists():
        failures.append("separate supplementary PDF")
    figure_files = [OUT / "figures" / f"Figure{i}.pdf" for i in range(1, 6)]
    if not all(p.exists() and p.stat().st_size > 0 for p in figure_files):
        failures.append("five separate vector PDF figure files")
    for required in [OUT / "cover_letter.docx",
                     OUT / "checklists" / "TRIPOD_AI_checklist.docx",
                     OUT / "checklists" / "TRIPOD_AI_abstract_checklist.docx"]:
        if not required.exists():
            failures.append(f"missing package file: {required.name}")

    paragraphs = doc.paragraphs
    ref_start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "References")
    ref_end = next(i for i, p in enumerate(paragraphs)
                   if p.text.strip() == "Acknowledgements and Funding")
    ref_paragraphs = [p for p in paragraphs[ref_start + 1:ref_end] if p.text.strip()]
    refs = list(range(1, len(ref_paragraphs) + 1))
    if not refs or len(refs) > 60 or any(p.style.name != "Compact" for p in ref_paragraphs):
        failures.append("sequential Nature-style reference list <=60")
    body = text[:text.find("References")]
    cited = cited_reference_numbers(body)
    if cited != set(refs):
        failures.append(f"citation closure (cited={sorted(cited)}, refs={refs})")

    if "[ARCHIVE DOI/URL REQUIRED BEFORE SUBMISSION]" in text:
        blockers.append("stable aggregate-archive DOI/URL is not inserted")
    commit_match = re.search(r"(?:at|at the)\s+(?:immutable\s+)?commit\s+`?([0-9a-f]{40})`?", text)
    if not commit_match:
        failures.append("immutable analysis commit in Code Availability")
    else:
        remote = subprocess.run(
            ["git", "branch", "-r", "--contains", commit_match.group(1)],
            cwd=ROOT, text=True, capture_output=True, check=False).stdout
        if "srrepo/" not in remote:
            blockers.append("analysis commit is not pushed to the public Scientific Reports code repository")
    if re.search(r"\[(?:TODO|TBC|PLACEHOLDER|INSERT)[^]]*\]", text, re.I):
        failures.append("no unresolved placeholders")

    cover_path = OUT / "cover_letter.docx"
    if cover_path.exists():
        cover_text = document_text(Document(cover_path))
        if "[AUTHOR CONFIRMATION REQUIRED]" in cover_text or "The authors must confirm" in cover_text:
            blockers.append(
                "all-author approval and suggested-reviewer conflict confirmations are pending"
            )

    report = {"status": "FAIL" if failures else ("BLOCKED" if blockers else "PASS"),
              "failures": failures, "submission_blockers": blockers,
              "main_docx_bytes": MAIN.stat().st_size,
              "main_tables": len(doc.tables), "figure_legends": len(legends)}
    (OUT / "verification_report.json").write_text(
        json.dumps(report, indent=2), encoding="utf-8")
    print(json.dumps(report, indent=2))
    return 1 if failures else (2 if blockers else 0)


if __name__ == "__main__":
    sys.exit(main())
