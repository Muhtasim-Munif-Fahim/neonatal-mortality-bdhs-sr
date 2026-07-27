"""Build the corrected Scientific Reports manuscript and supplement.

All numerical tokens in the manuscript template are replaced from regenerated
machine-readable outputs. The main DOCX contains five editable legends and
three editable tables, but no embedded main figures.
"""
from __future__ import annotations

import json
import re
import subprocess
import zipfile
from xml.etree import ElementTree as ET
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results"
FIGURES = ROOT / "figures"
OUT = ROOT / "manuscript" / "scientific_reports"
TEMPLATE = OUT / "manuscript_scientific_reports_template.md"
STUDIO = OUT / "figure_studio"
TITLE = ("Temporal evaluation of neonatal mortality prediction and decomposition "
         "of mortality decline in Bangladesh, 2011–2022")


def analysis_commit() -> str:
    return subprocess.check_output(["git", "rev-parse", "HEAD"], cwd=ROOT,
                                   text=True).strip()


def _fmt(value, digits=3, signed=False):
    return f"{float(value):+.{digits}f}" if signed else f"{float(value):.{digits}f}"


def _ci(row, lo="ci_low", hi="ci_high", digits=1, signed=False):
    return f"{_fmt(row[lo], digits, signed)} to {_fmt(row[hi], digits, signed)}"


def _cohort_row(cohort, name, year=None):
    sub = cohort[cohort.cohort.eq(name)]
    if year is not None:
        return sub[sub.survey_year.eq(year)].iloc[0]
    return pd.Series({"n": sub.n.sum(), "deaths": sub.deaths.sum()})


def _metadata() -> dict:
    path = OUT / "submission_metadata.json"
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return {"code_url": "[NEW CODE REPOSITORY URL REQUIRED]"}


def replacements() -> dict[str, str]:
    cohort = pd.read_csv(RESULTS / "cohort_definition.csv")
    lock = json.loads((RESULTS / "pipeline_lock.json").read_text())
    forward = pd.read_csv(RESULTS / "forward_validation.csv")
    summary = (forward.groupby("model", as_index=False)
               .agg(mean_PR_AUC=("mean_PR_AUC", "first"),
                    mean_Brier=("mean_Brier", "first"))
               .sort_values("mean_PR_AUC", ascending=False))
    second = summary[summary.model.ne(lock["model"])].iloc[0]
    perf = pd.read_csv(RESULTS / "primary_performance_uncertainty.csv").set_index("prediction")
    recal = pd.read_csv(RESULTS / "primary_recalibration.csv").set_index("prediction")
    scope = pd.read_csv(RESULTS / "table11_predictor_scope_sensitivity.csv")
    scope = scope[scope.prediction.eq("isotonic_recalibrated")].set_index("scope")
    native = pd.read_csv(RESULTS / "table10_survey_weighted_native.csv").set_index("training")
    trend = pd.read_csv(RESULTS / "table3_nmr_trend.csv").set_index("year")
    coverage = pd.read_csv(RESULTS / "table3b_prevalence.csv").set_index("exposure")
    dec = pd.read_csv(RESULTS / "table5_decomposition.csv").set_index("component")
    care_dec = pd.read_csv(RESULTS / "table5b_decomposition_care_module.csv").set_index("component")

    primary = _cohort_row(cohort, "primary_allrecent")
    dev_rows = cohort[(cohort.cohort.eq("primary_allrecent")) & cohort.survey_year.ne(2022)]
    dev = pd.Series({"n": dev_rows.n.sum(), "deaths": dev_rows.deaths.sum()})
    test = _cohort_row(cohort, "primary_allrecent", 2022)
    min1 = _cohort_row(cohort, "sensitivity_allrecent_min1")
    care = _cohort_row(cohort, "sensitivity_care_module")
    care_test = _cohort_row(cohort, "sensitivity_care_module", 2022)
    raw, cal = perf.loc["raw"], perf.loc["recalibrated"]
    cal_detail = recal.loc["isotonic_recalibrated"]
    context = scope.loc["survey_context_enriched"]
    care_scope = scope.loc["care_enriched_most_recent"]
    unweighted = native.loc["native_unweighted"]
    weighted = native.loc["native_survey_weighted"]

    values = {
        "N_PRIMARY": f"{int(primary.n):,}", "E_PRIMARY": f"{int(primary.deaths):,}",
        "N_DEV": f"{int(dev.n):,}", "E_DEV": f"{int(dev.deaths):,}",
        "N_TEST": f"{int(test.n):,}", "E_TEST": f"{int(test.deaths):,}",
        "N_MIN1": f"{int(min1.n):,}", "E_MIN1": f"{int(min1.deaths):,}",
        "N_CARE": f"{int(care.n):,}", "E_CARE": f"{int(care.deaths):,}",
        "N_CARE_TEST": f"{int(care_test.n):,}", "E_CARE_TEST": f"{int(care_test.deaths):,}",
        "LOCK_MODEL": str(lock["model"]),
        "LOCK_MEAN_AP": _fmt(lock["mean_pr_auc"]),
        "LOCK_MEAN_BRIER": _fmt(lock["mean_brier"], 5),
        "SECOND_MODEL": str(second.model),
        "SECOND_MEAN_AP": _fmt(second.mean_PR_AUC),
        "SECOND_MEAN_BRIER": _fmt(second.mean_Brier, 5),
        "RAW_AUROC": _fmt(raw.ROC_AUC), "RAW_AUROC_CI": str(raw.ROC_AUC_95CI),
        "RAW_AP": _fmt(raw.PR_AUC), "RAW_AP_CI": str(raw.PR_AUC_95CI),
        "RAW_BRIER": _fmt(raw.Brier, 5), "RAW_MEAN_RISK": _fmt(raw.mean_predicted_risk, 3),
        "CAL_AUROC": _fmt(cal.ROC_AUC), "CAL_AUROC_CI": str(cal.ROC_AUC_95CI),
        "CAL_AP": _fmt(cal.PR_AUC), "CAL_AP_CI": str(cal.PR_AUC_95CI),
        "CAL_BRIER": _fmt(cal.Brier, 5), "CAL_BRIER_CI": str(cal.Brier_95CI),
        "CAL_BRIER_SKILL": _fmt(cal_detail.Brier_skill, 3, signed=True),
        "CAL_MEAN_RISK": _fmt(cal.mean_predicted_risk, 3),
        "CAL_SLOPE": _fmt(cal.cal_slope), "CAL_SLOPE_CI": str(cal.cal_slope_95CI),
        "CAL_INTERCEPT": _fmt(cal.cal_intercept),
        "CAL_INTERCEPT_CI": str(cal.cal_intercept_95CI),
        "TEST_PREVALENCE": _fmt(cal.observed_prevalence, 3),
        "CONTEXT_AUROC": _fmt(context.ROC_AUC), "CONTEXT_AP": _fmt(context.PR_AUC),
        "CONTEXT_BRIER_SKILL": _fmt(context.Brier_skill, 3, signed=True),
        "CARE_AUROC": _fmt(care_scope.ROC_AUC), "CARE_AP": _fmt(care_scope.PR_AUC),
        "CARE_BRIER_SKILL": _fmt(care_scope.Brier_skill, 3, signed=True),
        "NATIVE_AUROC": _fmt(unweighted.ROC_AUC), "NATIVE_AP": _fmt(unweighted.PR_AUC),
        "NATIVE_BRIER": _fmt(unweighted.Brier, 5),
        "WEIGHTED_AUROC": _fmt(weighted.ROC_AUC), "WEIGHTED_AP": _fmt(weighted.PR_AUC),
        "NMR_2011": _fmt(trend.loc[2011].NMR, 1),
        "NMR_2011_CI": _ci(trend.loc[2011], digits=1),
        "NMR_2014": _fmt(trend.loc[2014].NMR, 1),
        "NMR_2017": _fmt(trend.loc[2017].NMR, 1),
        "NMR_2022": _fmt(trend.loc[2022].NMR, 1),
        "NMR_2022_CI": _ci(trend.loc[2022], digits=1),
        "FACILITY_2011": _fmt(coverage.loc["facility delivery", "y2011"], 1),
        "FACILITY_2022": _fmt(coverage.loc["facility delivery", "y2022"], 1),
        "SKILLED_2011": _fmt(coverage.loc["skilled attendant", "y2011"], 1),
        "SKILLED_2022": _fmt(coverage.loc["skilled attendant", "y2022"], 1),
        "CSECTION_2011": _fmt(coverage.loc["caesarean section", "y2011"], 1),
        "CSECTION_2022": _fmt(coverage.loc["caesarean section", "y2022"], 1),
        "DECOMP_NMR_2011": _fmt(dec.loc["NMR_2011"].value_per_1000, 1),
        "DECOMP_NMR_2022": _fmt(dec.loc["NMR_2022"].value_per_1000, 1),
        "DECOMP_TOTAL": _fmt(dec.loc["total_change"].value_per_1000, 1, True),
        "DECOMP_TOTAL_CI": _ci(dec.loc["total_change"], digits=1, signed=True),
        "DECOMP_DIST": _fmt(dec.loc["distribution"].value_per_1000, 1, True),
        "DECOMP_DIST_CI": _ci(dec.loc["distribution"], digits=1, signed=True),
        "DECOMP_EFFECT": _fmt(dec.loc["effect"].value_per_1000, 1, True),
        "DECOMP_EFFECT_CI": _ci(dec.loc["effect"], digits=1, signed=True),
        "CARE_DECOMP_TOTAL": _fmt(care_dec.loc["total_change"].value_per_1000, 1, True),
        "CARE_DECOMP_TOTAL_CI": _ci(care_dec.loc["total_change"], digits=1, signed=True),
        "CARE_DECOMP_DIST": _fmt(care_dec.loc["distribution"].value_per_1000, 1, True),
        "CARE_DECOMP_DIST_CI": _ci(care_dec.loc["distribution"], digits=1, signed=True),
        "CARE_DECOMP_EFFECT": _fmt(care_dec.loc["effect"].value_per_1000, 1, True),
        "CARE_DECOMP_EFFECT_CI": _ci(care_dec.loc["effect"], digits=1, signed=True),
        "ANALYSIS_COMMIT": analysis_commit(), "CODE_URL": _metadata()["code_url"],
    }
    legends = []
    for number in range(1, 6):
        legends.append((STUDIO / f"figure_{number}" / "caption.md").read_text(encoding="utf-8").strip())
    values["FIGURE_LEGENDS"] = "\n\n".join(legends)
    tables = []
    for number in range(1, 4):
        text = (OUT / "tables" / f"Table{number}.md").read_text(encoding="utf-8").strip()
        tables.append(text)
    values["MAIN_TABLES"] = "\n\n".join(tables)
    return values


def render_template() -> str:
    text = TEMPLATE.read_text(encoding="utf-8")
    for key, value in replacements().items():
        text = text.replace("{{" + key + "}}", str(value))
    unresolved = sorted(set(re.findall(r"\{\{([A-Z0-9_]+)\}\}", text)))
    if unresolved:
        raise RuntimeError(f"Unresolved manuscript tokens: {unresolved}")
    abstract = text.split("# Abstract", 1)[1].split("**Keywords:**", 1)[0]
    abstract_words = re.findall(r"\b[\w'-]+\b", abstract)
    if len(abstract_words) > 195:
        raise RuntimeError(f"Abstract has {len(abstract_words)} words (limit 195)")
    return text


def md_table(df: pd.DataFrame) -> str:
    frame = df.fillna("").astype(str)
    header = "| " + " | ".join(frame.columns) + " |"
    rule = "| " + " | ".join(["---"] * len(frame.columns)) + " |"
    rows = ["| " + " | ".join(value.replace("|", "\\|") for value in row) + " |"
            for row in frame.to_numpy().tolist()]
    return "\n".join([header, rule, *rows])


def supplement_markdown() -> str:
    sections = [
        "# Supplementary Information",
        f"## {TITLE}",
        "Supplementary analyses are sensitivity or exploratory analyses. They do not alter the development-only pipeline lock, and decision-curve or subgroup results do not establish clinical utility or fairness.",
    ]
    sources = [
        ("Supplementary Table S1. Outcome-stratified descriptive characteristics", "tableS1_outcome_descriptives.csv"),
        ("Supplementary Table S2. All candidate model results in 2022", "metrics_all.csv"),
        ("Supplementary Table S3. Forward chronological validation", "forward_validation.csv"),
        ("Supplementary Table S4. Balancing sensitivity", "table6_balancing.csv"),
        ("Supplementary Table S5. Candidate-model PSU-bootstrap intervals", "table7_bootstrap_ci.csv"),
        ("Supplementary Table S6. Exploratory subgroup performance", "table8_subgroups.csv"),
        ("Supplementary Table S7. Primary missing-indicator audit", "table9_indicator_sensitivity.csv"),
        ("Supplementary Table S8. Native-class and survey-weighted training sensitivity", "table10_survey_weighted_native.csv"),
        ("Supplementary Table S9. Maternity-service and composition trends", "table3b_prevalence.csv"),
        ("Supplementary Table S10. Association-shift analysis", "table4_association_shift.csv"),
        ("Supplementary Table S11. Care-module decomposition", "table5b_decomposition_care_module.csv"),
        ("Supplementary Table S12. Follow-up sensitivity and cohort arithmetic", "cohort_definition.csv"),
        ("Supplementary Table S13. Predictor-timing and questionnaire-scope sensitivity", "table11_predictor_scope_sensitivity.csv"),
        ("Supplementary Table S14. Adjusted associations with neonatal death", "table_adjusted_or.csv"),
    ]
    for title, filename in sources:
        path = RESULTS / filename
        if not path.exists():
            continue
        frame = pd.read_csv(path)
        for column in frame.select_dtypes(include="number").columns:
            frame[column] = frame[column].round(4)
        sections.extend([f"## {title}", md_table(frame)])
    images = [
        ("Supplementary Figure S1. Boruta feature-selection comparison.", "fig2_boruta.png"),
        ("Supplementary Figure S2. Exploratory decision-curve analysis.", "figS_decision_curve.png"),
        ("Supplementary Figure S3. Balancing and recalibration sensitivity.", "fig10_balancing_calibration.png"),
        ("Supplementary Figure S4. SHAP attribution shares across survey rounds.", "fig6_shap_temporal.png"),
    ]
    for title, filename in images:
        path = FIGURES / filename
        if path.exists():
            sections.extend([f"## {title}", f"![]({path.as_posix()}){{width=6.2in}}"])
    return "\n\n".join(sections) + "\n"


def strip_empty_comments_part(path: Path) -> None:
    temporary = path.with_suffix(".without-comments.docx")
    with zipfile.ZipFile(path) as source, zipfile.ZipFile(
            temporary, "w", compression=zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            if item.filename == "word/comments.xml":
                continue
            payload = source.read(item.filename)
            if item.filename == "[Content_Types].xml":
                root = ET.fromstring(payload)
                for node in list(root):
                    if node.attrib.get("PartName") == "/word/comments.xml":
                        root.remove(node)
                payload = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif item.filename == "word/_rels/document.xml.rels":
                root = ET.fromstring(payload)
                for node in list(root):
                    if node.attrib.get("Type", "").endswith("/comments"):
                        root.remove(node)
                payload = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            target.writestr(item, payload)
    temporary.replace(path)


def _field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    code = OxmlElement("w:instrText"); code.set(qn("xml:space"), "preserve"); code.text = instruction
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for node in (begin, code, separate, text, end):
        run._r.append(node)


def _style_font(style, name: str, size: float, bold: bool = False):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = None
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)


def _table_widths(table, total_dxa: int) -> list[int]:
    columns = len(table.columns)
    weights = []
    for column in range(columns):
        longest = max(len(cell.text.strip()) for cell in table.columns[column].cells)
        weights.append(max(5.0, min(30.0, longest ** 0.5 * 3.0)))
    scale = total_dxa / sum(weights)
    widths = [max(420, round(weight * scale)) for weight in weights]
    widths[-1] += total_dxa - sum(widths)
    return widths


def _set_table_geometry(table, total_dxa: int, font_size: float):
    table.autofit = False
    widths = _table_widths(table, total_dxa)
    properties = table._tbl.tblPr
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders"); properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}"); borders.append(node)
        node.set(qn("w:val"), "single"); node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), "B7BCC5")
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout"); properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW"); properties.append(table_width)
    table_width.set(qn("w:type"), "dxa"); table_width.set(qn("w:w"), str(total_dxa))
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd"); properties.append(indent)
    indent.set(qn("w:type"), "dxa"); indent.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol"); column.set(qn("w:w"), str(width)); grid.append(column)
    for row_number, row in enumerate(table.rows):
        row_properties = row._tr.get_or_add_trPr()
        if row_number == 0:
            repeat = OxmlElement("w:tblHeader"); repeat.set(qn("w:val"), "true")
            row_properties.append(repeat)
        for column_number, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(widths[column_number]))
            if row_number == 0:
                shading = OxmlElement("w:shd"); shading.set(qn("w:fill"), "F2F4F7")
                cell._tc.get_or_add_tcPr().append(shading)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.alignment = (WD_ALIGN_PARAGRAPH.LEFT if column_number == 0
                                       else WD_ALIGN_PARAGRAPH.CENTER)
                for run in paragraph.runs:
                    run.font.name = "Arial"; run.font.size = Pt(font_size)
                    run.bold = row_number == 0
                    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Arial")
                    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Arial")


def patch_docx(path: Path):
    document = Document(path)
    supplement = path.name.startswith("supplementary")
    for section in document.sections:
        if supplement:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = Inches(11), Inches(8.5)
            section.top_margin = section.bottom_margin = Inches(0.55)
            section.left_margin = section.right_margin = Inches(0.55)
        else:
            section.page_width, section.page_height = Inches(8.5), Inches(11)
            section.top_margin = section.bottom_margin = Inches(1.0)
            section.left_margin = section.right_margin = Inches(1.0)
        line_numbers = section._sectPr.find(qn("w:lnNumType"))
        if line_numbers is None:
            line_numbers = OxmlElement("w:lnNumType")
            section._sectPr.append(line_numbers)
        line_numbers.set(qn("w:countBy"), "1")
        line_numbers.set(qn("w:restart"), "newPage")
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not footer.text:
            _field(footer, "PAGE")
    normal = document.styles["Normal"]
    _style_font(normal, "Arial", 10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)
    for style_name, size in [("Title", 16), ("Heading 1", 13), ("Heading 2", 11)]:
        if style_name in document.styles:
            style = document.styles[style_name]
            _style_font(style, "Arial", size, True)
            style.paragraph_format.keep_with_next = True
            style.paragraph_format.space_before = Pt(12 if style_name != "Title" else 0)
            style.paragraph_format.space_after = Pt(6)
    for table in document.tables:
        _set_table_geometry(table, 14256 if supplement else 9360,
                            6.2 if supplement else 7.2)
    settings = document.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields"); settings.append(update)
    update.set(qn("w:val"), "false")
    document.core_properties.title = TITLE
    document.core_properties.author = "Md. Muhtasim Munif Fahim; Md. Naim Molla; Md. Rezaul Karim"
    document.core_properties.last_modified_by = "Authors"
    document.save(path)
    strip_empty_comments_part(path)


def build():
    OUT.mkdir(parents=True, exist_ok=True)
    main_md = OUT / "manuscript_scientific_reports.md"
    supplement_md = OUT / "supplementary_information.md"
    main_md.write_text(render_template(), encoding="utf-8")
    supplement_md.write_text(supplement_markdown(), encoding="utf-8")
    for source, target in [
        (main_md, OUT / "manuscript_scientific_reports.docx"),
        (supplement_md, OUT / "supplementary_information.docx"),
    ]:
        subprocess.run(["pandoc", str(source), "-o", str(target), "--standalone"], check=True)
        patch_docx(target)
    print(OUT / "manuscript_scientific_reports.docx")
    print(OUT / "supplementary_information.docx")


if __name__ == "__main__":
    build()
